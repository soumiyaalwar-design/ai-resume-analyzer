"""
AI Resume Analyzer (Advanced Edition)
--------------------------------------
A Streamlit web app that analyzes resumes (PDF/DOCX), extracts skills using
NLP (spaCy), checks contact info, sections, ATS-friendliness, action verbs,
experience level, named entities, and compares against job roles or pasted
job descriptions. Includes visual charts and score history tracking.

Author: AI Resume Analyzer Team
"""

import re
import io
from datetime import datetime

import streamlit as st
import PyPDF2
import spacy
from spacy.matcher import PhraseMatcher
import matplotlib.pyplot as plt
from docx import Document


# ----------------------------------------------------------------------
# CONFIGURATION & DATA
# ----------------------------------------------------------------------

SKILL_DB = [
    "python", "java", "c++", "c", "javascript", "sql", "r",
    "machine learning", "deep learning", "nlp", "natural language processing",
    "computer vision", "data analysis", "data visualization",
    "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "keras",
    "matplotlib", "seaborn", "power bi", "tableau", "excel",
    "django", "flask", "fastapi", "streamlit",
    "html", "css", "react", "node.js", "git", "docker", "kubernetes",
    "aws", "azure", "gcp", "linux", "mongodb", "mysql", "postgresql",
    "statistics", "data structures", "algorithms", "rest api",
    "communication", "teamwork", "problem solving", "leadership"
]

# Skill categories for pie chart (feature 9)
SKILL_CATEGORIES = {
    "Programming": ["python", "java", "c++", "c", "javascript", "sql", "r", "html", "css"],
    "ML / AI": ["machine learning", "deep learning", "nlp", "natural language processing",
                "computer vision", "scikit-learn", "tensorflow", "pytorch", "keras"],
    "Data & Visualization": ["pandas", "numpy", "data analysis", "data visualization",
                              "power bi", "tableau", "excel", "matplotlib", "seaborn", "statistics"],
    "Tools & Frameworks": ["django", "flask", "fastapi", "streamlit", "react", "node.js",
                            "git", "docker", "kubernetes", "aws", "azure", "gcp", "linux",
                            "mongodb", "mysql", "postgresql", "rest api"],
    "CS Fundamentals": ["data structures", "algorithms"],
    "Soft Skills": ["communication", "teamwork", "problem solving", "leadership"],
}

JOB_ROLES = {
    "Data Scientist": [
        "python", "sql", "machine learning", "deep learning",
        "pandas", "numpy", "scikit-learn", "data visualization",
        "statistics", "tensorflow"
    ],
    "AI Engineer": [
        "python", "machine learning", "deep learning", "nlp",
        "tensorflow", "pytorch", "computer vision", "docker",
        "git", "algorithms"
    ],
    "Data Analyst": [
        "sql", "excel", "power bi", "tableau", "python",
        "data visualization", "statistics", "data analysis"
    ],
    "Machine Learning Engineer": [
        "python", "machine learning", "deep learning", "tensorflow",
        "pytorch", "scikit-learn", "docker", "kubernetes", "git",
        "rest api"
    ],
    "Web Developer": [
        "html", "css", "javascript", "react", "node.js",
        "django", "flask", "git", "mongodb", "rest api"
    ],
    "Software Engineer": [
        "python", "java", "c++", "data structures", "algorithms",
        "git", "sql", "rest api", "linux", "problem solving"
    ]
}

# Resume sections we expect to find (feature 2 groundwork, used in ATS check)
EXPECTED_SECTIONS = {
    "Education": ["education", "academic"],
    "Experience": ["experience", "work history", "employment"],
    "Skills": ["skills", "technical skills", "core competencies"],
    "Projects": ["projects", "personal projects"],
    "Certifications": ["certification", "certificate", "courses"],
}

# Action verbs (feature 4 groundwork, used in ATS check)
STRONG_VERBS = [
    "developed", "designed", "implemented", "built", "created", "led",
    "managed", "optimized", "improved", "increased", "reduced", "automated",
    "deployed", "architected", "launched", "analyzed", "achieved",
    "delivered", "streamlined", "spearheaded", "collaborated", "mentored",
    "engineered", "established", "executed", "initiated"
]

WEAK_PHRASES = [
    "responsible for", "worked on", "helped with", "involved in",
    "duties included", "tasked with"
]

# Common degree abbreviations / keywords for NER-style extraction (feature 14)
DEGREE_PATTERNS = [
    r"b\.?tech", r"m\.?tech", r"b\.?e\.?", r"m\.?e\.?", r"b\.?sc",
    r"m\.?sc", r"mba", r"mca", r"bca", r"bachelor[^.,\n]*", r"master[^.,\n]*",
    r"ph\.?d", r"diploma[^.,\n]*"
]


# ----------------------------------------------------------------------
# NLP MODEL LOADING
# ----------------------------------------------------------------------

@st.cache_resource
def load_spacy_model():
    """Load the spaCy English language model (cached)."""
    try:
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        from spacy.cli import download
        download("en_core_web_sm")
        nlp = spacy.load("en_core_web_sm")
    return nlp


@st.cache_resource
def build_phrase_matcher(_nlp):
    """Build a spaCy PhraseMatcher loaded with all skills from SKILL_DB."""
    matcher = PhraseMatcher(_nlp.vocab, attr="LOWER")
    patterns = [_nlp.make_doc(skill) for skill in SKILL_DB]
    matcher.add("SKILLS", patterns)
    return matcher


# ----------------------------------------------------------------------
# FILE TEXT EXTRACTION (PDF + DOCX) - FEATURE 10
# ----------------------------------------------------------------------

def extract_text_from_pdf(uploaded_file):
    """Extract raw text from an uploaded PDF file using PyPDF2."""
    text = ""
    try:
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        st.error(f"Error reading PDF file: {e}")
    return text


def extract_text_from_docx(uploaded_file):
    """Extract raw text from an uploaded DOCX file using python-docx."""
    text = ""
    try:
        doc = Document(io.BytesIO(uploaded_file.read()))
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        # Also extract text from tables (common in resumes)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + " "
                text += "\n"
    except Exception as e:
        st.error(f"Error reading DOCX file: {e}")
    return text


def extract_text_from_file(uploaded_file):
    """
    Dispatch text extraction based on file extension.

    Args:
        uploaded_file: Streamlit UploadedFile object (.pdf or .docx)

    Returns:
        str: Extracted raw text
    """
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    else:
        st.error("Unsupported file format. Please upload a PDF or DOCX file.")
        return ""


# ----------------------------------------------------------------------
# TEXT CLEANING / PREPROCESSING
# ----------------------------------------------------------------------

def clean_text(text):
    """Clean and normalize raw resume text."""
    text = text.lower()
    text = re.sub(r"[\n\r\t]", " ", text)
    text = re.sub(r"[^a-z0-9\+\#\.\s@]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


# ----------------------------------------------------------------------
# SKILL EXTRACTION (NLP-BASED USING SPACY PHRASE MATCHER)
# ----------------------------------------------------------------------

def extract_skills(text, nlp, matcher):
    """Extract skills from cleaned text using spaCy's PhraseMatcher."""
    doc = nlp(text)
    matches = matcher(doc)

    found_skills = set()
    for match_id, start, end in matches:
        span = doc[start:end]
        found_skills.add(span.text.lower().strip())

    return found_skills


# ----------------------------------------------------------------------
# RESUME SCORE CALCULATION
# ----------------------------------------------------------------------

def calculate_resume_score(found_skills):
    """Calculate an overall resume score (out of 100)."""
    if not SKILL_DB:
        return 0
    score = (len(found_skills) / len(SKILL_DB)) * 100
    return round(min(score, 100))


def score_label(score):
    """Return a (label, color) tuple based on score range."""
    if score >= 70:
        return "Excellent", "#059669"
    elif score >= 40:
        return "Good", "#d97706"
    else:
        return "Needs Work", "#dc2626"


# ----------------------------------------------------------------------
# JOB ROLE MATCHING
# ----------------------------------------------------------------------

def match_skills(found_skills, required_skills):
    """
    Generic matcher: compare found skills against any set of required skills.

    Returns:
        tuple: (matched_skills, missing_skills, match_percentage)
    """
    required_skills = set(required_skills)
    matched_skills = found_skills.intersection(required_skills)
    missing_skills = required_skills.difference(found_skills)

    if not required_skills:
        match_percentage = 0
    else:
        match_percentage = (len(matched_skills) / len(required_skills)) * 100

    return matched_skills, missing_skills, round(match_percentage, 1)


def compare_all_roles(found_skills):
    """
    Compare resume skills against every job role in JOB_ROLES (Feature 6).

    Returns:
        dict: {role_name: match_percentage}
    """
    results = {}
    for role, required in JOB_ROLES.items():
        _, _, pct = match_skills(found_skills, required)
        results[role] = pct
    return results


# ----------------------------------------------------------------------
# JOB DESCRIPTION SKILL EXTRACTION - FEATURES 8 & 13
# ----------------------------------------------------------------------

def extract_skills_from_jd(jd_text, nlp, matcher):
    """
    Extract required skills directly from a pasted job description,
    using the same skill matcher as the resume.

    Args:
        jd_text (str): Raw job description text
        nlp: spaCy model
        matcher: PhraseMatcher

    Returns:
        set: Skills mentioned in the JD
    """
    cleaned = clean_text(jd_text)
    return extract_skills(cleaned, nlp, matcher)


# ----------------------------------------------------------------------
# CONTACT INFO CHECKER - FEATURE 5
# ----------------------------------------------------------------------

def check_contact_info(raw_text):
    """
    Check for presence of email, phone number, LinkedIn, and GitHub links.

    Args:
        raw_text (str): Raw (uncleaned) resume text, to preserve symbols

    Returns:
        dict: {field_name: bool}
    """
    text_lower = raw_text.lower()

    email_found = bool(re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", raw_text))
    phone_found = bool(re.search(r"(\+?\d{1,3}[-.\s]?)?\(?\d{3,5}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}", raw_text))
    linkedin_found = "linkedin.com" in text_lower or "linkedin" in text_lower
    github_found = "github.com" in text_lower or "github" in text_lower

    return {
        "Email": email_found,
        "Phone Number": phone_found,
        "LinkedIn": linkedin_found,
        "GitHub": github_found,
    }


# ----------------------------------------------------------------------
# RESUME SECTION DETECTOR (used in ATS check, feature 11)
# ----------------------------------------------------------------------

def detect_sections(text_lower):
    """
    Check whether standard resume sections appear in the text.

    Args:
        text_lower (str): Lowercased resume text

    Returns:
        dict: {section_name: bool}
    """
    found_sections = {}
    for section, keywords in EXPECTED_SECTIONS.items():
        found_sections[section] = any(kw in text_lower for kw in keywords)
    return found_sections


# ----------------------------------------------------------------------
# EXPERIENCE LEVEL ESTIMATION - FEATURE GROUNDWORK FOR ATS/SUMMARY
# ----------------------------------------------------------------------

def estimate_experience(text_lower):
    """
    Estimate years of experience and seniority level from resume text.

    Looks for patterns like "3+ years", "5 years of experience", or
    date ranges like "2019 - 2023" / "2021 - present".

    Args:
        text_lower (str): Lowercased resume text

    Returns:
        tuple: (estimated_years (int or None), level (str))
    """
    years_found = []

    # Pattern 1: "X years" or "X+ years"
    for match in re.findall(r"(\d+)\+?\s*(?:years|yrs)", text_lower):
        years_found.append(int(match))

    # Pattern 2: date ranges like "2019 - 2023" or "2020 - present"
    current_year = datetime.now().year
    for start, end in re.findall(r"(20\d{2})\s*[-–to]+\s*(20\d{2}|present|current)", text_lower):
        start_year = int(start)
        end_year = current_year if end in ("present", "current") else int(end)
        if end_year >= start_year:
            years_found.append(end_year - start_year)

    if not years_found:
        return None, "Entry Level"

    total_years = max(years_found)

    if total_years <= 1:
        level = "Entry Level"
    elif total_years <= 4:
        level = "Mid Level"
    else:
        level = "Senior Level"

    return total_years, level


# ----------------------------------------------------------------------
# ACTION VERB ANALYZER (used in ATS check)
# ----------------------------------------------------------------------

def analyze_action_verbs(text_lower):
    """
    Count strong action verbs vs weak passive phrases.

    Returns:
        dict: {"strong_count": int, "weak_count": int,
               "strong_found": list, "weak_found": list}
    """
    strong_found = [v for v in STRONG_VERBS if v in text_lower]
    weak_found = [p for p in WEAK_PHRASES if p in text_lower]

    return {
        "strong_count": len(strong_found),
        "weak_count": len(weak_found),
        "strong_found": strong_found,
        "weak_found": weak_found,
    }


# ----------------------------------------------------------------------
# ATS-FRIENDLINESS CHECK - FEATURE 11
# ----------------------------------------------------------------------

def run_ats_check(raw_text, cleaned_text, found_skills):
    """
    Run a basic simulated ATS (Applicant Tracking System) check.

    Returns:
        list of dicts: [{"check": str, "passed": bool, "detail": str}, ...]
    """
    text_lower = raw_text.lower()
    word_count = len(cleaned_text.split())
    results = []

    # 1. Resume length check
    if word_count < 150:
        results.append({"check": "Resume Length", "passed": False,
                         "detail": f"Only {word_count} words. Aim for 300-700 words for adequate detail."})
    elif word_count > 1200:
        results.append({"check": "Resume Length", "passed": False,
                         "detail": f"{word_count} words is quite long. Consider trimming to 1-2 pages."})
    else:
        results.append({"check": "Resume Length", "passed": True,
                         "detail": f"{word_count} words — a healthy length."})

    # 2. Skills section presence
    sections = detect_sections(text_lower)
    if sections["Skills"]:
        results.append({"check": "Skills Section", "passed": True,
                         "detail": "A dedicated skills section was found."})
    else:
        results.append({"check": "Skills Section", "passed": False,
                         "detail": "No clear 'Skills' section header found. ATS systems often look for this explicitly."})

    # 3. Contact info
    contact = check_contact_info(raw_text)
    missing_contact = [k for k, v in contact.items() if not v]
    if not missing_contact:
        results.append({"check": "Contact Information", "passed": True,
                         "detail": "Email, phone, LinkedIn, and GitHub all detected."})
    else:
        results.append({"check": "Contact Information", "passed": False,
                         "detail": f"Missing: {', '.join(missing_contact)}."})

    # 4. Action verbs vs weak phrases
    verbs = analyze_action_verbs(text_lower)
    if verbs["weak_count"] > verbs["strong_count"]:
        results.append({"check": "Action Verbs", "passed": False,
                         "detail": f"Found {verbs['weak_count']} weak/passive phrases vs "
                                   f"{verbs['strong_count']} strong action verbs. "
                                   "Replace phrases like 'responsible for' with action verbs."})
    else:
        results.append({"check": "Action Verbs", "passed": True,
                         "detail": f"Good use of action verbs ({verbs['strong_count']} found)."})

    # 5. Special character density (tables/columns often break ATS parsing)
    special_char_ratio = len(re.findall(r"[^\w\s]", raw_text)) / max(len(raw_text), 1)
    if special_char_ratio > 0.05:
        results.append({"check": "Formatting Complexity", "passed": False,
                         "detail": "High density of special characters/symbols detected. "
                                   "Complex tables, columns, or icons can confuse ATS parsers. "
                                   "Prefer simple single-column layouts."})
    else:
        results.append({"check": "Formatting Complexity", "passed": True,
                         "detail": "Text structure looks simple and ATS-friendly."})

    # 6. Skill count check
    if len(found_skills) < 5:
        results.append({"check": "Skill Coverage", "passed": False,
                         "detail": f"Only {len(found_skills)} recognizable skills found. "
                                   "Add more relevant tools, languages, and frameworks."})
    else:
        results.append({"check": "Skill Coverage", "passed": True,
                         "detail": f"{len(found_skills)} relevant skills detected."})

    return results


# ----------------------------------------------------------------------
# NAMED ENTITY EXTRACTION - FEATURE 14
# ----------------------------------------------------------------------

def extract_entities(raw_text, nlp):
    """
    Extract organizations (companies/universities) using spaCy NER,
    and degrees using regex patterns.

    Args:
        raw_text (str): Raw resume text
        nlp: spaCy model

    Returns:
        dict: {"organizations": list, "degrees": list}
    """
    # Limit text length for performance on large resumes
    doc = nlp(raw_text[:100000])

    orgs = []
    for ent in doc.ents:
        if ent.label_ == "ORG":
            name = ent.text.strip()
            # Filter out very short or noisy matches
            if 2 < len(name) < 60 and name not in orgs:
                orgs.append(name)

    degrees = []
    text_lower = raw_text.lower()
    for pattern in DEGREE_PATTERNS:
        for match in re.findall(pattern, text_lower):
            match_clean = match.strip().title()
            if match_clean not in degrees:
                degrees.append(match_clean)

    return {
        "organizations": orgs[:10],
        "degrees": degrees[:5],
    }


# ----------------------------------------------------------------------
# AI-STYLE SUMMARY GENERATOR - FEATURE 12 (RULE-BASED)
# ----------------------------------------------------------------------

def generate_ai_summary(resume_score, match_percentage, job_role, found_skills,
                         matched_skills, missing_skills, experience_level,
                         contact_info, ats_results):
    """
    Generate a short, personalized natural-language summary of the resume's
    strengths and weaknesses, similar to an AI reviewer's verdict.

    This is rule-based (no external API) so it works offline and free,
    but reads like an automated AI assessment.

    Returns:
        str: 3-5 sentence summary
    """
    sentences = []

    # Opening line based on overall score
    label, _ = score_label(resume_score)
    sentences.append(
        f"This resume scores **{resume_score}/100** overall, which is rated "
        f"**{label.lower()}** based on detected technical and soft skills."
    )

    # Job match line
    if match_percentage >= 70:
        sentences.append(
            f"It aligns strongly with the **{job_role}** role, matching "
            f"{len(matched_skills)} of the required skills."
        )
    elif match_percentage >= 40:
        sentences.append(
            f"It shows a moderate fit for **{job_role}**, but is missing "
            f"{len(missing_skills)} key skills that recruiters often screen for."
        )
    else:
        sentences.append(
            f"It currently shows a weak fit for **{job_role}**, with "
            f"{len(missing_skills)} important skills not detected."
        )

    # Experience line
    sentences.append(f"Based on the content, the candidate appears to be at the **{experience_level}**.")

    # ATS issues line
    failed_checks = [r["check"] for r in ats_results if not r["passed"]]
    if failed_checks:
        sentences.append(
            f"A few areas may reduce ATS compatibility: {', '.join(failed_checks)}."
        )
    else:
        sentences.append("The resume passes all basic ATS-friendliness checks.")

    # Contact info line
    missing_contact = [k for k, v in contact_info.items() if not v]
    if missing_contact:
        sentences.append(f"Consider adding missing contact details: {', '.join(missing_contact)}.")

    return " ".join(sentences)


# ----------------------------------------------------------------------
# IMPROVEMENT SUGGESTIONS
# ----------------------------------------------------------------------

def generate_suggestions(missing_skills, match_percentage, resume_score,
                          ats_results, contact_info, verbs_analysis):
    """Generate human-readable improvement suggestions."""
    suggestions = []

    if missing_skills:
        skill_list = ", ".join(sorted(s.title() for s in missing_skills))
        suggestions.append(
            f"Consider adding or highlighting these skills relevant to "
            f"your target role: **{skill_list}**."
        )

    if match_percentage < 50:
        suggestions.append(
            "Your resume matches less than half of the required skills "
            "for this role. Consider taking relevant courses or projects "
            "to strengthen your profile."
        )
    elif match_percentage < 80:
        suggestions.append(
            "You're a decent match for this role! Adding the missing "
            "skills above could significantly boost your chances."
        )
    else:
        suggestions.append(
            "Great job! Your resume aligns strongly with this role's "
            "requirements."
        )

    if resume_score < 30:
        suggestions.append(
            "Your overall skill coverage is low. Try adding more "
            "technical skills, tools, and certifications relevant to "
            "your field."
        )

    # Contact info suggestions
    missing_contact = [k for k, v in contact_info.items() if not v]
    if missing_contact:
        suggestions.append(
            f"Add the following missing contact details so recruiters can "
            f"reach you: **{', '.join(missing_contact)}**."
        )

    # Action verbs suggestion
    if verbs_analysis["weak_count"] > verbs_analysis["strong_count"]:
        suggestions.append(
            "Replace passive phrases like 'responsible for' or 'worked on' "
            "with strong action verbs such as 'Developed', 'Led', or "
            "'Optimized' to make your impact clearer."
        )

    # ATS-based suggestions
    for check in ats_results:
        if not check["passed"]:
            suggestions.append(f"**{check['check']}**: {check['detail']}")

    suggestions.append(
        "Quantify your achievements with numbers (e.g., 'Improved model "
        "accuracy by 15%') to make your resume more impactful."
    )

    return suggestions


# ----------------------------------------------------------------------
# REPORT GENERATION
# ----------------------------------------------------------------------

def generate_report_text(job_role, resume_score, match_percentage, found_skills,
                          matched_skills, missing_skills, suggestions, ai_summary):
    """Generate a plain-text summary report for download."""
    lines = [
        "AI RESUME ANALYZER - REPORT",
        "=" * 40,
        f"Target Job Role: {job_role}",
        f"Overall Resume Score: {resume_score}%",
        f"Job Match Percentage: {match_percentage}%",
        "",
        "AI Summary:",
        ai_summary.replace("**", ""),
        "",
        "Extracted Skills:",
        ", ".join(sorted(s.title() for s in found_skills)) or "None",
        "",
        "Matched Skills:",
        ", ".join(sorted(s.title() for s in matched_skills)) or "None",
        "",
        "Missing Skills:",
        ", ".join(sorted(s.title() for s in missing_skills)) or "None",
        "",
        "Improvement Suggestions:",
    ]
    for i, s in enumerate(suggestions, 1):
        clean_s = s.replace("**", "")
        lines.append(f"{i}. {clean_s}")

    return "\n".join(lines)


# ----------------------------------------------------------------------
# CHARTS - FEATURE 9
# ----------------------------------------------------------------------

def plot_skill_category_pie(found_skills):
    """
    Create a pie chart showing distribution of found skills across
    SKILL_CATEGORIES.

    Returns:
        matplotlib.figure.Figure or None if no skills found
    """
    category_counts = {}
    for category, skills in SKILL_CATEGORIES.items():
        count = len(found_skills.intersection(set(skills)))
        if count > 0:
            category_counts[category] = count

    if not category_counts:
        return None

    fig, ax = plt.subplots(figsize=(4.5, 4.5))
    colors = ["#4f46e5", "#2563eb", "#059669", "#d97706", "#dc2626", "#7c3aed"]
    ax.pie(
        category_counts.values(),
        labels=category_counts.keys(),
        autopct="%1.0f%%",
        colors=colors[:len(category_counts)],
        textprops={"fontsize": 9}
    )
    ax.set_title("Skill Category Distribution", fontsize=11, fontweight="bold")
    fig.tight_layout()
    return fig


def plot_matched_vs_missing(matched_count, missing_count, role_label):
    """
    Create a horizontal bar chart comparing matched vs missing skills
    for the selected job role.

    Returns:
        matplotlib.figure.Figure
    """
    fig, ax = plt.subplots(figsize=(4.5, 2.2))
    categories = ["Matched", "Missing"]
    values = [matched_count, missing_count]
    colors = ["#059669", "#dc2626"]

    bars = ax.barh(categories, values, color=colors)
    ax.set_title(f"Skill Match: {role_label}", fontsize=11, fontweight="bold")
    ax.set_xlabel("Number of skills")

    for bar, value in zip(bars, values):
        ax.text(bar.get_width() + 0.1, bar.get_y() + bar.get_height() / 2,
                str(value), va="center", fontsize=10)

    fig.tight_layout()
    return fig


def plot_role_comparison(role_scores):
    """
    Create a horizontal bar chart comparing match % across all job roles.

    Args:
        role_scores (dict): {role_name: match_percentage}

    Returns:
        matplotlib.figure.Figure
    """
    sorted_roles = dict(sorted(role_scores.items(), key=lambda x: x[1]))

    fig, ax = plt.subplots(figsize=(6, 3.5))
    bars = ax.barh(list(sorted_roles.keys()), list(sorted_roles.values()), color="#4f46e5")
    ax.set_xlim(0, 100)
    ax.set_xlabel("Match Percentage (%)")
    ax.set_title("Your Fit Across Different Job Roles", fontsize=11, fontweight="bold")

    for bar, value in zip(bars, sorted_roles.values()):
        ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height() / 2,
                f"{value}%", va="center", fontsize=9)

    fig.tight_layout()
    return fig


# ----------------------------------------------------------------------
# STREAMLIT UI - CUSTOM STYLING
# ----------------------------------------------------------------------

def load_custom_css():
    """Inject custom CSS for a modern, polished look."""
    st.markdown(
        """
        <style>
        .block-container {
            padding-top: 2rem;
            padding-bottom: 3rem;
            max-width: 1100px;
        }

        .header-banner {
            background: linear-gradient(135deg, #4f46e5 0%, #7c3aed 50%, #2563eb 100%);
            padding: 40px 35px;
            border-radius: 16px;
            margin-bottom: 30px;
            box-shadow: 0 4px 20px rgba(79, 70, 229, 0.25);
        }

        .header-banner h1 {
            color: #ffffff;
            font-size: 38px;
            font-weight: 800;
            margin: 0;
        }

        .header-banner p {
            color: rgba(255,255,255,0.9);
            font-size: 16px;
            margin-top: 8px;
            margin-bottom: 0;
        }

        .metric-card {
            background-color: #ffffff;
            padding: 24px 20px;
            border-radius: 16px;
            box-shadow: 0 4px 14px rgba(0,0,0,0.07);
            text-align: center;
            border: 1px solid #eef0f5;
            transition: transform 0.15s ease;
        }
        .metric-card:hover {
            transform: translateY(-3px);
            box-shadow: 0 8px 24px rgba(0,0,0,0.10);
        }

        .metric-label {
            font-size: 13px;
            color: #6b7280;
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            margin-bottom: 6px;
        }

        .metric-value {
            font-size: 32px;
            font-weight: 800;
            color: #1f2a44;
            margin: 4px 0;
        }

        .skill-pill {
            display: inline-block;
            background: linear-gradient(135deg, #dbeafe, #e0e7ff);
            color: #1e40af;
            padding: 7px 16px;
            margin: 4px;
            border-radius: 24px;
            font-size: 13px;
            font-weight: 600;
            border: 1px solid #c7d2fe;
        }

        .matched-pill {
            display: inline-block;
            background: linear-gradient(135deg, #d1fae5, #ccfbf1);
            color: #047857;
            padding: 7px 16px;
            margin: 4px;
            border-radius: 24px;
            font-size: 13px;
            font-weight: 600;
            border: 1px solid #a7f3d0;
        }

        .missing-pill {
            display: inline-block;
            background: linear-gradient(135deg, #fee2e2, #fde8e8);
            color: #b91c1c;
            padding: 7px 16px;
            margin: 4px;
            border-radius: 24px;
            font-size: 13px;
            font-weight: 600;
            border: 1px solid #fecaca;
        }

        .suggestion-box {
            background-color: #fffbeb;
            border-left: 4px solid #f59e0b;
            padding: 14px 18px;
            border-radius: 10px;
            margin-bottom: 12px;
            font-size: 14.5px;
            color: #374151;
            box-shadow: 0 2px 8px rgba(0,0,0,0.04);
        }

        .section-header {
            font-size: 22px;
            font-weight: 700;
            color: #1f2a44;
            margin-top: 10px;
            margin-bottom: 14px;
        }

        .summary-box {
            background: linear-gradient(135deg, #eef2ff, #f5f3ff);
            border: 1px solid #e0e7ff;
            border-radius: 14px;
            padding: 18px 22px;
            font-size: 15px;
            line-height: 1.7;
            color: #1f2a44;
            margin-bottom: 10px;
        }

        .check-pass {
            color: #059669;
            font-weight: 700;
        }
        .check-fail {
            color: #dc2626;
            font-weight: 700;
        }

        .info-pill {
            display: inline-block;
            background: #f3f4f6;
            color: #374151;
            padding: 6px 14px;
            margin: 4px;
            border-radius: 20px;
            font-size: 13px;
            font-weight: 600;
            border: 1px solid #e5e7eb;
        }

        .stProgress > div > div > div > div {
            background: linear-gradient(90deg, #4f46e5, #2563eb);
        }

        section[data-testid="stSidebar"] {
            background-color: #f8f9fc;
            border-right: 1px solid #e5e7eb;
        }

        .empty-state {
            text-align:center;
            padding: 60px 20px;
            background:#f8f9fc;
            border-radius:16px;
            border: 2px dashed #c7d2fe;
        }
        .empty-state h2 {
            color:#4f46e5;
            margin-bottom:8px;
        }
        .empty-state p {
            color:#6b7280;
            font-size:15px;
            line-height: 1.6;
        }

        footer {visibility: hidden;}
        #MainMenu {visibility: hidden;}
        </style>
        """,
        unsafe_allow_html=True
    )


# ----------------------------------------------------------------------
# UI HELPER COMPONENTS
# ----------------------------------------------------------------------

def render_skill_pills(skills, pill_class="skill-pill"):
    """Render a set of skills as styled HTML 'pills'."""
    if not skills:
        st.info("None found.")
        return

    pills_html = "".join(
        f'<span class="{pill_class}">{skill.title()}</span>'
        for skill in sorted(skills)
    )
    st.markdown(pills_html, unsafe_allow_html=True)


def render_metric_card(label, value, suffix="%", show_label=False):
    """Render a styled metric card with a label, large value, and progress bar."""
    extra_html = ""
    if show_label and suffix == "%":
        text, color = score_label(value)
        extra_html = (
            f'<div style="color:{color}; font-weight:700; '
            f'font-size:13px; margin-top:4px;">{text}</div>'
        )

    st.markdown(
        f"""
        <div class="metric-card">
            <div class="metric-label">{label}</div>
            <div class="metric-value">{value}{suffix}</div>
            {extra_html}
        </div>
        """,
        unsafe_allow_html=True
    )
    if suffix == "%":
        st.progress(min(int(value), 100) / 100)


def render_ats_results(ats_results):
    """Render ATS check results as a clean checklist."""
    for check in ats_results:
        icon = "✅" if check["passed"] else "⚠️"
        css_class = "check-pass" if check["passed"] else "check-fail"
        st.markdown(
            f"{icon} <span class='{css_class}'>{check['check']}</span> — {check['detail']}",
            unsafe_allow_html=True
        )


def render_info_pills(items):
    """Render a list of plain info pills (used for organizations/degrees)."""
    if not items:
        st.info("None detected.")
        return
    pills_html = "".join(f'<span class="info-pill">{item}</span>' for item in items)
    st.markdown(pills_html, unsafe_allow_html=True)


# ----------------------------------------------------------------------
# MAIN APPLICATION
# ----------------------------------------------------------------------

def main():
    st.set_page_config(
        page_title="AI Resume Analyzer",
        page_icon="📄",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    load_custom_css()

    # ---------------- HEADER ----------------
    st.markdown(
        """
        <div class="header-banner">
            <h1>📄 AI Resume Analyzer</h1>
            <p>Upload your resume (PDF or DOCX), select a target job role or paste a
            job description, and get an instant AI-powered analysis with charts,
            ATS checks, and personalized suggestions.</p>
        </div>
        """,
        unsafe_allow_html=True
    )

    # ---------------- SIDEBAR ----------------
    with st.sidebar:
        st.markdown("### 📄 AI Resume Analyzer")
        st.caption("Smart resume scoring & job-fit matching")
        st.markdown("---")

        job_role = st.selectbox(
            "🎯 Select Target Job Role",
            options=list(JOB_ROLES.keys()),
            help="Choose the role you're applying for to get a tailored match score."
        )

        st.markdown("---")
        st.markdown("#### 📝 Optional: Paste a Job Description")
        jd_text = st.text_area(
            "Paste a job description to match against instead of (or alongside) the role above.",
            height=120,
            placeholder="Paste the full job description here..."
        )

        st.markdown("---")
        st.markdown("#### 📂 Upload Resume")
        uploaded_file = st.file_uploader(
            "Upload your resume (PDF or DOCX)",
            type=["pdf", "docx"],
            help="PDF and DOCX files are supported."
        )

        if uploaded_file:
            st.success(f"✅ {uploaded_file.name}")
            if st.button("🔄 Analyze Another Resume", use_container_width=True):
                st.session_state.pop("analysis", None)
                st.session_state.pop("file_id", None)
                st.session_state.pop("job_role", None)
                st.session_state.pop("jd_text", None)
                st.rerun()

        st.markdown("---")
        st.markdown(
            "**ℹ️ How it works**\n\n"
            "1. Upload your resume (PDF/DOCX)\n"
            "2. Select a job role and/or paste a job description\n"
            "3. Get scores, charts, ATS checks, and suggestions instantly"
        )

    # ---------------- LOAD NLP MODEL ----------------
    nlp = load_spacy_model()
    matcher = build_phrase_matcher(nlp)

    # ---------------- EMPTY STATE ----------------
    if uploaded_file is None:
        st.markdown(
            """
            <div class="empty-state">
                <h2>👋 Welcome!</h2>
                <p>
                    Upload your resume (PDF or DOCX) from the sidebar to get started.<br>
                    We'll extract your skills, score your resume, check ATS-friendliness,
                    and match it against your target job role or a pasted job description.
                </p>
            </div>
            """,
            unsafe_allow_html=True
        )
        return

    # ---------------- ANALYSIS (WITH SESSION STATE CACHING) ----------------
    file_id = getattr(uploaded_file, "file_id", uploaded_file.name)

    needs_analysis = (
        "analysis" not in st.session_state
        or st.session_state.get("file_id") != file_id
        or st.session_state.get("job_role") != job_role
        or st.session_state.get("jd_text") != jd_text
    )

    if needs_analysis:
        status = st.status("Analyzing your resume...", expanded=True)

        status.write("📖 Extracting text from file...")
        raw_text = extract_text_from_file(uploaded_file)

        if not raw_text.strip():
            status.update(label="Extraction failed", state="error")
            st.error(
                "Could not extract any text from this file. "
                "Please make sure it's a text-based PDF/DOCX, not a scanned image."
            )
            return

        status.write("🧹 Cleaning text...")
        cleaned_text = clean_text(raw_text)
        text_lower = raw_text.lower()

        status.write("🧠 Extracting skills with NLP...")
        found_skills = extract_skills(cleaned_text, nlp, matcher)

        status.write("📇 Checking contact information...")
        contact_info = check_contact_info(raw_text)

        status.write("📂 Detecting resume sections & experience level...")
        sections = detect_sections(text_lower)
        exp_years, exp_level = estimate_experience(text_lower)
        verbs_analysis = analyze_action_verbs(text_lower)

        status.write("🔍 Running ATS-friendliness checks...")
        ats_results = run_ats_check(raw_text, cleaned_text, found_skills)

        status.write("🏢 Extracting organizations & education info...")
        entities = extract_entities(raw_text, nlp)

        status.write("📊 Calculating scores...")
        resume_score = calculate_resume_score(found_skills)

        # Determine which skill set to match against: JD (if provided) or job role
        if jd_text and jd_text.strip():
            jd_skills = extract_skills_from_jd(jd_text, nlp, matcher)
            matched_skills, missing_skills, match_percentage = match_skills(found_skills, jd_skills)
            match_label = "Job Description"
        else:
            required = JOB_ROLES.get(job_role, [])
            matched_skills, missing_skills, match_percentage = match_skills(found_skills, required)
            match_label = job_role

        status.write("📈 Comparing against all job roles...")
        role_scores = compare_all_roles(found_skills)

        status.write("✍️ Generating AI summary...")
        ai_summary = generate_ai_summary(
            resume_score, match_percentage, match_label, found_skills,
            matched_skills, missing_skills, exp_level, contact_info, ats_results
        )

        suggestions = generate_suggestions(
            missing_skills, match_percentage, resume_score,
            ats_results, contact_info, verbs_analysis
        )

        status.update(label="✅ Analysis complete!", state="complete", expanded=False)

        analysis_data = {
            "raw_text": raw_text,
            "found_skills": found_skills,
            "resume_score": resume_score,
            "matched_skills": matched_skills,
            "missing_skills": missing_skills,
            "match_percentage": match_percentage,
            "match_label": match_label,
            "suggestions": suggestions,
            "contact_info": contact_info,
            "sections": sections,
            "exp_years": exp_years,
            "exp_level": exp_level,
            "verbs_analysis": verbs_analysis,
            "ats_results": ats_results,
            "entities": entities,
            "role_scores": role_scores,
            "ai_summary": ai_summary,
        }

        st.session_state["analysis"] = analysis_data
        st.session_state["file_id"] = file_id
        st.session_state["job_role"] = job_role
        st.session_state["jd_text"] = jd_text

        # ---- FEATURE 15: Score history tracking ----
        if "history" not in st.session_state:
            st.session_state["history"] = []

        st.session_state["history"].append({
            "filename": uploaded_file.name,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "resume_score": resume_score,
            "match_percentage": match_percentage,
            "match_label": match_label,
        })

    data = st.session_state["analysis"]
    raw_text = data["raw_text"]
    found_skills = data["found_skills"]
    resume_score = data["resume_score"]
    matched_skills = data["matched_skills"]
    missing_skills = data["missing_skills"]
    match_percentage = data["match_percentage"]
    match_label = data["match_label"]
    suggestions = data["suggestions"]
    contact_info = data["contact_info"]
    sections = data["sections"]
    exp_years = data["exp_years"]
    exp_level = data["exp_level"]
    ats_results = data["ats_results"]
    entities = data["entities"]
    role_scores = data["role_scores"]
    ai_summary = data["ai_summary"]

    # ================== AI SUMMARY (Feature 12) ==================
    st.markdown('<div class="section-header">🤖 AI Summary</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="summary-box">{ai_summary}</div>', unsafe_allow_html=True)

    # ================== SCORE OVERVIEW ==================
    st.markdown('<div class="section-header">📊 Score Overview</div>', unsafe_allow_html=True)

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        render_metric_card("Overall Resume Score", resume_score, show_label=True)
    with col2:
        render_metric_card(f"Match: {match_label}", match_percentage, show_label=True)
    with col3:
        render_metric_card("Skills Detected", len(found_skills), suffix="")
    with col4:
        exp_display = f"{exp_years}+ yrs" if exp_years else exp_level
        render_metric_card("Experience", exp_display, suffix="")

    # ================== EXTRACTED SKILLS ==================
    st.markdown('<div class="section-header">🛠️ Extracted Skills</div>', unsafe_allow_html=True)
    render_skill_pills(found_skills, "skill-pill")

    # ================== JOB / JD MATCH ==================
    st.markdown(
        f'<div class="section-header">🎯 Match for {match_label}</div>',
        unsafe_allow_html=True
    )

    col_a, col_b = st.columns(2)
    with col_a:
        st.markdown("**✅ Matched Skills**")
        render_skill_pills(matched_skills, "matched-pill")
    with col_b:
        st.markdown("**❌ Missing Skills**")
        render_skill_pills(missing_skills, "missing-pill")

    # ================== CHARTS (Feature 9) ==================
    st.markdown('<div class="section-header">📈 Visual Insights</div>', unsafe_allow_html=True)

    chart_col1, chart_col2 = st.columns(2)
    with chart_col1:
        pie_fig = plot_skill_category_pie(found_skills)
        if pie_fig:
            st.pyplot(pie_fig, use_container_width=True)
        else:
            st.info("Not enough categorized skills to display a chart.")
    with chart_col2:
        bar_fig = plot_matched_vs_missing(len(matched_skills), len(missing_skills), match_label)
        st.pyplot(bar_fig, use_container_width=True)

    # ================== JOB ROLE COMPARISON (Feature 6) ==================
    st.markdown('<div class="section-header">🧭 Fit Across Job Roles</div>', unsafe_allow_html=True)
    role_fig = plot_role_comparison(role_scores)
    st.pyplot(role_fig, use_container_width=True)

    best_role = max(role_scores, key=role_scores.get)
    st.markdown(
        f"💡 Based on your current skills, you appear to be the **strongest fit for "
        f"{best_role}** ({role_scores[best_role]}% match)."
    )

    # ================== CONTACT INFO & SECTIONS (Feature 5) ==================
    st.markdown('<div class="section-header">📇 Resume Completeness Check</div>', unsafe_allow_html=True)

    col_c, col_d = st.columns(2)
    with col_c:
        st.markdown("**Contact Information**")
        for field, present in contact_info.items():
            icon = "✅" if present else "❌"
            st.markdown(f"{icon} {field}")

    with col_d:
        st.markdown("**Resume Sections**")
        for section, present in sections.items():
            icon = "✅" if present else "❌"
            st.markdown(f"{icon} {section}")

    # ================== EDUCATION & ORGANIZATIONS (Feature 14) ==================
    st.markdown('<div class="section-header">🎓 Detected Education & Organizations</div>', unsafe_allow_html=True)

    col_e, col_f = st.columns(2)
    with col_e:
        st.markdown("**Degrees / Qualifications**")
        render_info_pills(entities["degrees"])
    with col_f:
        st.markdown("**Organizations Mentioned**")
        render_info_pills(entities["organizations"])

    # ================== ATS CHECK (Feature 11) ==================
    st.markdown('<div class="section-header">🔍 ATS-Friendliness Check</div>', unsafe_allow_html=True)
    render_ats_results(ats_results)

    # ================== SUGGESTIONS ==================
    st.markdown('<div class="section-header">💡 Improvement Suggestions</div>', unsafe_allow_html=True)
    for suggestion in suggestions:
        st.markdown(f'<div class="suggestion-box">{suggestion}</div>', unsafe_allow_html=True)

    # ================== SCORE HISTORY (Feature 15) ==================
    history = st.session_state.get("history", [])
    if len(history) > 1:
        st.markdown('<div class="section-header">🕒 Score History (This Session)</div>', unsafe_allow_html=True)
        st.table([
            {
                "File": h["filename"],
                "Time": h["timestamp"],
                "Resume Score": f"{h['resume_score']}%",
                "Match": f"{h['match_percentage']}% ({h['match_label']})",
            }
            for h in history
        ])

        if len(history) >= 2:
            prev = history[-2]
            curr = history[-1]
            score_diff = curr["resume_score"] - prev["resume_score"]
            if score_diff > 0:
                st.success(f"📈 Your resume score improved by {score_diff} points compared to your last upload!")
            elif score_diff < 0:
                st.warning(f"📉 Your resume score dropped by {abs(score_diff)} points compared to your last upload.")
            else:
                st.info("Your resume score is unchanged compared to your last upload.")

    # ================== DOWNLOAD REPORT ==================
    st.markdown("<br>", unsafe_allow_html=True)
    report_text = generate_report_text(
        match_label, resume_score, match_percentage,
        found_skills, matched_skills, missing_skills, suggestions, ai_summary
    )
    st.download_button(
        label="⬇️ Download Full Report (.txt)",
        data=report_text,
        file_name="resume_analysis_report.txt",
        mime="text/plain"
    )

    # ================== RAW TEXT (EXPANDABLE) ==================
    with st.expander("📄 View Extracted Resume Text"):
        st.text_area("Extracted Text", raw_text, height=300)


if __name__ == "__main__":
    main()